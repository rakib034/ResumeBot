#for docx
'''import os
import streamlit as st
from docx import Document
from dotenv import load_dotenv
from groq import Groq

# Load Groq API key
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
groq_client = Groq(api_key=GROQ_API_KEY)

# Folder paths
UPLOAD_DIR = "uploads"
DOWNLOAD_DIR = "download"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

st.set_page_config(page_title="Resume Editor (Preserve Format)", layout="wide")
st.title("📄 Resume Editor with AI (Formatting Preserved)")

# Select file
doc_files = [f for f in os.listdir(UPLOAD_DIR) if f.endswith(".docx")]
selected_file = st.sidebar.selectbox("Select a resume", doc_files)

if selected_file:
    doc_path = os.path.join(UPLOAD_DIR, selected_file)
    doc = Document(doc_path)

    st.subheader("📝 Edit Resume Content")
    updated_texts = []

    for i, para in enumerate(doc.paragraphs):
        key = f"para_{i}"
        text = para.text.strip()
        new_text = st.text_input(f"Paragraph {i+1}", value=text, key=key)
        updated_texts.append(new_text)

    # --- AI Chat Box ---
    st.sidebar.subheader("💬 AI Assistant")
    chat_query = st.sidebar.text_input("Ask something about resume writing")

    if chat_query:
        response = groq_client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[
                {"role": "system", "content": "You are a resume-writing assistant. Help the user write good resume content."},
                {"role": "user", "content": chat_query}
            ]
        )
        suggestion = response.choices[0].message.content
        st.sidebar.markdown("**AI Suggestion:**")
        st.sidebar.code(suggestion)

    # --- Save button ---
    output_name = st.text_input("Enter filename to save:", value=f"edited_{selected_file}")
    if st.button("💾 Save Edited Resume"):
        for i, new_text in enumerate(updated_texts):
            doc.paragraphs[i].text = new_text  # Update only text, formatting preserved

        save_path = os.path.join(DOWNLOAD_DIR, output_name)
        doc.save(save_path)
        st.success(f"Saved to: {save_path}")

        with open(save_path, "rb") as f:
            st.download_button("📥 Download File", f, file_name=output_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
'''





#bottom code needs docx2html to convert docx to html for preview
'''import os
import streamlit as st
from docx import Document
from dotenv import load_dotenv
from groq import Groq
from docx2html import convert  # <-- Add this import

# Load Groq API key
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
groq_client = Groq(api_key=GROQ_API_KEY)

# Folder paths
UPLOAD_DIR = "uploads"
DOWNLOAD_DIR = "download"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

st.set_page_config(page_title="Resume Editor (Preserve Format)", layout="wide")
st.title("📄 Resume Editor with AI (Formatting Preserved)")

# Select file
doc_files = [f for f in os.listdir(UPLOAD_DIR) if f.endswith(".docx")]
selected_file = st.sidebar.selectbox("Select a resume", doc_files)

if selected_file:
    doc_path = os.path.join(UPLOAD_DIR, selected_file)
    doc = Document(doc_path)

    # Session state for edit mode
    if "edit_mode" not in st.session_state:
        st.session_state.edit_mode = False

    if not st.session_state.edit_mode:
        st.subheader("👀 Preview Original Resume (with formatting)")
        html = convert(doc_path)
        st.markdown(html, unsafe_allow_html=True)
        if st.button("✏️ Edit Resume"):
            st.session_state.edit_mode = True
    else:
        st.subheader("📝 Edit Resume Content")
        updated_texts = []

        for i, para in enumerate(doc.paragraphs):
            key = f"para_{i}"
            text = para.text.strip()
            new_text = st.text_input(f"Paragraph {i+1}", value=text, key=key)
            updated_texts.append(new_text)

        # --- AI Chat Box ---
        st.sidebar.subheader("💬 AI Assistant")
        chat_query = st.sidebar.text_input("Ask something about resume writing")

        if chat_query:
            response = groq_client.chat.completions.create(
                model="llama3-70b-8192",
                messages=[
                    {"role": "system", "content": "You are a resume-writing assistant. Help the user write good resume content."},
                    {"role": "user", "content": chat_query}
                ]
            )
            suggestion = response.choices[0].message.content
            st.sidebar.markdown("**AI Suggestion:**")
            st.sidebar.code(suggestion)

        # --- Save button ---
        output_name = st.text_input("Enter filename to save:", value=f"edited_{selected_file}")
        if st.button("💾 Save Edited Resume"):
            for i, new_text in enumerate(updated_texts):
                doc.paragraphs[i].text = new_text  # Update only text, formatting preserved

            save_path = os.path.join(DOWNLOAD_DIR, output_name)
            doc.save(save_path)
            st.success(f"Saved to: {save_path}")

            with open(save_path, "rb") as f:
                st.download_button("📥 Download File", f, file_name=output_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        if st.button("🔙 Back to Preview"):
            st.session_state.edit_mode = False'''